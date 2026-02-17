import unittest

from docx import Document
from docx.oxml.ns import qn

from docx_editor import add_table, delete_block, delete_content, delete_table, delete_text

_W_P = qn("w:p")
_W_TBL = qn("w:tbl")


def _flow_types(document):
    # Keep track of body order (paragraph/table) for flow-sensitive deletes.
    flow = []
    for element in document._body._element.iterchildren():
        if element.tag == _W_P:
            flow.append("paragraph")
        elif element.tag == _W_TBL:
            flow.append("table")
    return flow


class TestAddTable(unittest.TestCase):
    def test_add_table_infers_shape_from_data(self):
        document = Document()
        table = add_table(document, data=[["A", "B", "C"], ["D"]])

        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(table.rows), 2)
        self.assertEqual(len(table.columns), 3)
        self.assertEqual(table.cell(0, 0).text, "A")
        self.assertEqual(table.cell(0, 2).text, "C")
        self.assertEqual(table.cell(1, 0).text, "D")

    def test_add_table_can_insert_into_body_flow(self):
        document = Document()
        document.add_paragraph("Intro")
        document.add_paragraph("Outro")

        table = add_table(document, rows=1, cols=2, position=1)
        table.cell(0, 0).text = "X"

        self.assertEqual(_flow_types(document), ["paragraph", "table", "paragraph"])
        self.assertEqual(document.tables[0].cell(0, 0).text, "X")


class TestDeleteTable(unittest.TestCase):
    def test_delete_table_removes_structure_and_content(self):
        document = Document()
        document.add_paragraph("before")
        table = add_table(document, rows=1, cols=1)
        table.cell(0, 0).text = "to delete"
        document.add_paragraph("after")

        removed = delete_table(document, table_index=0)

        self.assertEqual(removed, 1)
        self.assertEqual(len(document.tables), 0)
        self.assertEqual(_flow_types(document), ["paragraph", "paragraph"])

    def test_delete_table_by_text_supports_remove_all(self):
        document = Document()
        first = add_table(document, rows=1, cols=1)
        second = add_table(document, rows=1, cols=1)
        first.cell(0, 0).text = "remove me"
        second.cell(0, 0).text = "remove me too"

        removed = delete_table(document, contains_text="remove me", remove_all=True)

        self.assertEqual(removed, 2)
        self.assertEqual(len(document.tables), 0)


class TestDeleteTextAndFlow(unittest.TestCase):
    def test_delete_text_handles_paragraphs_and_tables(self):
        document = Document()
        document.add_paragraph("alpha remove alpha remove")
        table = add_table(document, rows=1, cols=1)
        table.cell(0, 0).text = "remove in table"

        replaced = delete_text(document, "remove")

        self.assertEqual(replaced, 3)
        self.assertNotIn("remove", document.paragraphs[0].text.lower())
        self.assertNotIn("remove", table.cell(0, 0).text.lower())

    def test_delete_block_removes_table_without_breaking_flow(self):
        document = Document()
        document.add_paragraph("first")
        add_table(document, rows=1, cols=1)
        document.add_paragraph("last")

        removed_type = delete_block(document, 1)

        self.assertEqual(removed_type, "table")
        self.assertEqual(_flow_types(document), ["paragraph", "paragraph"])
        self.assertEqual([p.text for p in document.paragraphs], ["first", "last"])

    def test_delete_content_unified_selector(self):
        document = Document()
        document.add_paragraph("keep")
        table = add_table(document, rows=1, cols=1)
        table.cell(0, 0).text = "target"

        removed = delete_content(document, table_contains_text="target")

        self.assertEqual(removed, 1)
        self.assertEqual(len(document.tables), 0)


if __name__ == "__main__":
    unittest.main()
