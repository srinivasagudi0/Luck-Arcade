"""Small helper utilities for basic DOCX table/text editing."""

from __future__ import annotations

import re
from collections.abc import Sequence
from pathlib import Path
from typing import Any

from docx import Document as _DocumentFactory
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

_W_P = qn("w:p")
_W_TBL = qn("w:tbl")


def create_document() -> DocxDocument:
    """Make a new blank DOCX document."""
    return _DocumentFactory()


def open_document(path: str | Path) -> DocxDocument:
    """Open a DOCX file from disk."""
    return _DocumentFactory(str(path))


def save_document(document: DocxDocument, path: str | Path) -> None:
    """Save a DOCX file to disk."""
    document.save(str(path))


def _body_block_elements(document: DocxDocument) -> list[Any]:
    return [
        element
        for element in document._body._element.iterchildren()
        if element.tag in {_W_P, _W_TBL}
    ]


def _normalize_insert_index(index: int, length: int) -> int:
    if index < 0:
        index += length
    if index < 0:
        return 0
    if index > length:
        return length
    return index


def _normalize_index(index: int, length: int) -> int:
    if length == 0:
        raise IndexError("No items available to index.")
    if index < 0:
        index += length
    if index < 0 or index >= length:
        raise IndexError(f"Index {index} is out of range for {length} items.")
    return index


def add_table(
    document: DocxDocument,
    *,
    rows: int | None = None,
    cols: int | None = None,
    data: Sequence[Sequence[Any]] | None = None,
    position: int | None = None,
    style: str | None = None,
) -> Table:
    """Add a table, either by fixed size or inferred from `data`."""
    if rows is not None and rows < 0:
        raise ValueError("rows must be >= 0")
    if cols is not None and cols < 0:
        raise ValueError("cols must be >= 0")

    rows_data = [list(row) for row in data] if data else []
    inferred_rows = len(rows_data)
    inferred_cols = max((len(row) for row in rows_data), default=0)

    target_rows = max(rows or 0, inferred_rows)
    target_cols = max(cols or 0, inferred_cols)
    if target_rows <= 0 or target_cols <= 0:
        raise ValueError("Need rows/cols or non-empty data to size the table.")

    existing_blocks = _body_block_elements(document)

    table = document.add_table(rows=target_rows, cols=target_cols)
    if style:
        table.style = style

    for row_idx, row_values in enumerate(rows_data):
        for col_idx, value in enumerate(row_values):
            table.cell(row_idx, col_idx).text = "" if value is None else str(value)

    if position is not None:
        target_position = _normalize_insert_index(position, len(existing_blocks))
        if target_position < len(existing_blocks):
            existing_blocks[target_position].addprevious(table._tbl)

    return table


def _remove_table(table: Table) -> None:
    table_xml = table._tbl
    parent = table_xml.getparent()
    if parent is not None:
        parent.remove(table_xml)


def _table_text(table: Table) -> str:
    return "\n".join(cell.text for row in table.rows for cell in row.cells)


def delete_table(
    document: DocxDocument,
    *,
    table_index: int | None = None,
    contains_text: str | None = None,
    case_sensitive: bool = False,
    remove_all: bool = False,
) -> int:
    """Delete table(s) by index or by a text match."""
    by_index = table_index is not None
    by_text = contains_text is not None
    if by_index == by_text:
        raise ValueError("Specify exactly one of table_index or contains_text.")

    if by_index:
        tables = list(document.tables)
        index = _normalize_index(table_index, len(tables))
        _remove_table(tables[index])
        return 1

    assert contains_text is not None
    needle = contains_text if case_sensitive else contains_text.lower()
    matches: list[Table] = []
    for table in list(document.tables):
        haystack = _table_text(table)
        if not case_sensitive:
            haystack = haystack.lower()
        if needle in haystack:
            matches.append(table)
            if not remove_all:
                break

    for table in matches:
        _remove_table(table)
    return len(matches)


def _iter_paragraphs(document: DocxDocument, include_tables: bool) -> list[Paragraph]:
    paragraphs = list(document.paragraphs)
    if not include_tables:
        return paragraphs

    seen = {id(paragraph) for paragraph in paragraphs}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph_id = id(paragraph)
                    if paragraph_id in seen:
                        continue
                    seen.add(paragraph_id)
                    paragraphs.append(paragraph)
    return paragraphs


def _replace_in_paragraph(paragraph: Paragraph, pattern: re.Pattern[str]) -> int:
    if not paragraph.runs:
        updated, replaced = pattern.subn("", paragraph.text)
        if replaced:
            paragraph.text = updated
        return replaced

    total = 0
    for run in paragraph.runs:
        if not run.text:
            continue
        updated, replaced = pattern.subn("", run.text)
        if replaced:
            run.text = updated
            total += replaced
    return total


def delete_text(
    document: DocxDocument,
    text: str,
    *,
    case_sensitive: bool = False,
    whole_word: bool = False,
    include_tables: bool = True,
) -> int:
    """Delete matching text from paragraphs and optionally table cells."""
    if not text:
        return 0

    pattern_text = rf"\b{re.escape(text)}\b" if whole_word else re.escape(text)
    flags = 0 if case_sensitive else re.IGNORECASE
    pattern = re.compile(pattern_text, flags)

    replacements = 0
    for paragraph in _iter_paragraphs(document, include_tables=include_tables):
        replacements += _replace_in_paragraph(paragraph, pattern)
    return replacements


def delete_block(document: DocxDocument, block_index: int) -> str:
    """Delete a paragraph or table by body-flow index."""
    blocks = _body_block_elements(document)
    index = _normalize_index(block_index, len(blocks))
    block = blocks[index]
    block_type = "table" if block.tag == _W_TBL else "paragraph"
    parent = block.getparent()
    if parent is not None:
        parent.remove(block)
    return block_type


def delete_content(
    document: DocxDocument,
    *,
    text: str | None = None,
    table_index: int | None = None,
    table_contains_text: str | None = None,
    block_index: int | None = None,
    case_sensitive: bool = False,
    whole_word: bool = False,
    include_tables_for_text: bool = True,
    remove_all_matching_tables: bool = False,
) -> int:
    """Delete text, table, or a single flow block using one selector."""
    selectors = [
        text is not None,
        table_index is not None,
        table_contains_text is not None,
        block_index is not None,
    ]
    if sum(selectors) != 1:
        raise ValueError("Specify exactly one deletion target.")

    if text is not None:
        return delete_text(
            document,
            text,
            case_sensitive=case_sensitive,
            whole_word=whole_word,
            include_tables=include_tables_for_text,
        )
    if table_index is not None:
        return delete_table(document, table_index=table_index)
    if table_contains_text is not None:
        return delete_table(
            document,
            contains_text=table_contains_text,
            case_sensitive=case_sensitive,
            remove_all=remove_all_matching_tables,
        )

    assert block_index is not None
    delete_block(document, block_index)
    return 1


__all__ = [
    "add_table",
    "create_document",
    "delete_block",
    "delete_content",
    "delete_table",
    "delete_text",
    "open_document",
    "save_document",
]
